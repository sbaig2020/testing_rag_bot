import os
import re
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from datetime import datetime

import PyPDF2
from docx import Document
import markdown
from bs4 import BeautifulSoup
import json
import csv

from config_simple import settings

# Configure logging
logging.basicConfig(level=getattr(logging, settings.log_level))
logger = logging.getLogger(__name__)

@dataclass
class DocumentChunk:
    """Represents a chunk of processed document"""
    content: str
    metadata: Dict[str, Any]
    chunk_id: str
    source_file: str
    chunk_index: int

class DocumentProcessor:
    """Handles document processing and text extraction"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.txt': self._process_txt,
            '.docx': self._process_docx,
            '.md': self._process_markdown,
            '.html': self._process_html,
            '.csv': self._process_csv,
            '.json': self._process_json
        }
    
    def process_document(self, file_path: str) -> List[DocumentChunk]:
        """Process a document and return chunks"""
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            extension = file_path.suffix.lower()
            if extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {extension}")
            
            # Extract text using appropriate processor
            processor = self.supported_formats[extension]
            text_content, metadata = processor(file_path)
            
            # Create chunks
            chunks = self._create_chunks(text_content, str(file_path), metadata)
            
            logger.info(f"Processed {file_path.name}: {len(chunks)} chunks created")
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            raise
    
    def _process_pdf(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Extract text from PDF file"""
        text = ""
        metadata = {"file_type": "pdf", "pages": 0}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["pages"] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}"
                    
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
            raise
        
        return text, metadata
    
    def _process_txt(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            metadata = {
                "file_type": "txt",
                "char_count": len(text),
                "line_count": len(text.splitlines())
            }
            
            return text, metadata
            
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
            
            metadata = {
                "file_type": "txt",
                "encoding": "latin-1",
                "char_count": len(text),
                "line_count": len(text.splitlines())
            }
            
            return text, metadata
    
    def _process_docx(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            metadata = {
                "file_type": "docx",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables)
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            raise
    
    def _process_markdown(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Extract text from Markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            # Convert markdown to HTML then extract text
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            text = soup.get_text()
            
            # Also keep the original markdown for context
            text = f"Original Markdown:\n{md_content}\n\nPlain Text:\n{text}"
            
            metadata = {
                "file_type": "markdown",
                "original_length": len(md_content),
                "processed_length": len(text)
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error processing Markdown {file_path}: {str(e)}")
            raise
    
    def _process_html(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Extract text from HTML file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            metadata = {
                "file_type": "html",
                "title": soup.title.string if soup.title else "No title",
                "links": len(soup.find_all('a')),
                "images": len(soup.find_all('img'))
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error processing HTML {file_path}: {str(e)}")
            raise
    
    def _process_csv(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Extract text from CSV file"""
        try:
            text = ""
            row_count = 0
            
            with open(file_path, 'r', encoding='utf-8') as file:
                csv_reader = csv.reader(file)
                headers = next(csv_reader, None)
                
                if headers:
                    text += f"Headers: {', '.join(headers)}\n\n"
                
                for row_num, row in enumerate(csv_reader):
                    text += f"Row {row_num + 1}: {', '.join(row)}\n"
                    row_count += 1
                    
                    # Limit rows to prevent huge files
                    if row_count > 1000:
                        text += f"\n... (truncated after 1000 rows)\n"
                        break
            
            metadata = {
                "file_type": "csv",
                "columns": len(headers) if headers else 0,
                "rows": row_count,
                "headers": headers
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error processing CSV {file_path}: {str(e)}")
            raise
    
    def _process_json(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Extract text from JSON file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                json_data = json.load(file)
            
            # Convert JSON to readable text
            text = json.dumps(json_data, indent=2, ensure_ascii=False)
            
            metadata = {
                "file_type": "json",
                "size": len(text),
                "structure": self._analyze_json_structure(json_data)
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error processing JSON {file_path}: {str(e)}")
            raise
    
    def _analyze_json_structure(self, data: Any) -> Dict[str, Any]:
        """Analyze JSON structure"""
        if isinstance(data, dict):
            return {
                "type": "object",
                "keys": len(data.keys()),
                "key_names": list(data.keys())[:10]  # First 10 keys
            }
        elif isinstance(data, list):
            return {
                "type": "array",
                "length": len(data),
                "item_types": list(set(type(item).__name__ for item in data[:10]))
            }
        else:
            return {"type": type(data).__name__}
    
    def _create_chunks(self, text: str, source_file: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Split text into chunks"""
        chunks = []
        
        # Clean and normalize text
        text = self._clean_text(text)
        
        # Split into chunks
        chunk_size = settings.max_chunk_size
        overlap = settings.chunk_overlap
        
        words = text.split()
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk_words = words[i:i + chunk_size]
            chunk_text = ' '.join(chunk_words)
            
            if len(chunk_text.strip()) < 50:  # Skip very small chunks
                continue
            
            chunk_metadata = {
                **metadata,
                "chunk_size": len(chunk_text),
                "word_count": len(chunk_words),
                "created_at": datetime.now().isoformat(),
                "source_file": source_file
            }
            
            chunk = DocumentChunk(
                content=chunk_text,
                metadata=chunk_metadata,
                chunk_id=f"{Path(source_file).stem}_{len(chunks)}",
                source_file=source_file,
                chunk_index=len(chunks)
            )
            
            chunks.append(chunk)
        
        return chunks
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)\[\]\{\}\"\'\/\\]', '', text)
        
        # Normalize line breaks
        text = re.sub(r'\n+', '\n', text)
        
        return text.strip()
    
    def get_document_info(self, file_path: str) -> Dict[str, Any]:
        """Get basic information about a document without processing"""
        try:
            file_path = Path(file_path)
            stat = file_path.stat()
            
            return {
                "filename": file_path.name,
                "extension": file_path.suffix.lower(),
                "size_bytes": stat.st_size,
                "size_mb": round(stat.st_size / (1024 * 1024), 2),
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "supported": file_path.suffix.lower() in self.supported_formats
            }
            
        except Exception as e:
            logger.error(f"Error getting document info for {file_path}: {str(e)}")
            return {"error": str(e)}
